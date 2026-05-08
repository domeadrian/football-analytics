"""
Report Document Part 3 - Romanian Liga I Deep Dive (Sections 24-29)
Covers: Squad Analysis, Player Scouting, Club Diagnostics, Transfer Strategy,
        Video / Match Event Analysis, Season Narrative & Predictions
"""

import os
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from report_data import IMG_DIR
from report_charts_4 import SHORT_NAMES, POS_GROUPS, IDEAL_SQUAD, _get_ro_players


def _add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the doc."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Light Shading Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)


def build_document_part3(doc, data):
    """Build the Romanian Liga I deep-dive sections (24-29)."""
    ro_standings = data["ro_standings"]
    ro_events = data["ro_events"]
    players = data["players"]
    playoff_teams_data = data.get("playoff_teams_data", [])
    playout_teams_data = data.get("playout_teams_data", [])

    ro_players = _get_ro_players(players)

    # =====================================================================
    # SECTION 24: ROMANIAN LIGA I - SQUAD ANALYSIS DEEP DIVE
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("24. Romanian Liga I - Squad Analysis Deep Dive", level=1)

    doc.add_paragraph(
        "This section provides an exhaustive squad-by-squad analysis of every Romanian Liga I club, "
        "covering age profiles, position coverage, nationality balance, and market value distribution. "
        "The analysis aims to identify structural strengths and weaknesses that drive on-pitch performance."
    )

    # 24.1 Squad Age Profile
    doc.add_heading("24.1 Squad Age Profile per Club", level=2)
    doc.add_paragraph(
        "The age profile of a squad reveals its development stage and renewal needs. "
        "Clubs with heavy concentration in the 30+ bracket face imminent squad turnover, "
        "while clubs dominated by U21 players may lack experience for title challenges."
    )
    if os.path.exists(f"{IMG_DIR}/ro1_squad_age_profile.png"):
        doc.add_picture(f"{IMG_DIR}/ro1_squad_age_profile.png", width=Inches(6))

    # Age analysis per team
    if not ro_players.empty:
        doc.add_heading("24.1.1 Key Age Findings", level=3)
        age_stats = ro_players.groupby("short_team")["age"].agg(["mean", "median", "min", "max", "count"])
        age_stats = age_stats.sort_values("mean")
        rows = []
        for team, r in age_stats.iterrows():
            young = len(ro_players[(ro_players["short_team"] == team) & (ro_players["age"] <= 21)])
            old = len(ro_players[(ro_players["short_team"] == team) & (ro_players["age"] >= 30)])
            rows.append([team, f"{r['mean']:.1f}", f"{r['median']:.1f}", int(r['min']),
                         int(r['max']), int(r['count']), young, old])
        _add_table(doc, ["Club", "Mean Age", "Median", "Min", "Max", "Squad Size", "U21", "30+"], rows)

        doc.add_paragraph("")
        # Interpretation
        youngest = age_stats.index[0]
        oldest = age_stats.index[-1]
        doc.add_paragraph(
            f"The youngest squad belongs to {youngest} (avg age {age_stats.loc[youngest, 'mean']:.1f}), "
            f"suggesting a development-oriented philosophy. Conversely, {oldest} "
            f"(avg {age_stats.loc[oldest, 'mean']:.1f}) relies heavily on experience and faces "
            f"significant renewal needs in the next 2-3 transfer windows."
        )

    # 24.2 Position Coverage
    doc.add_heading("24.2 Position Coverage Analysis", level=2)
    doc.add_paragraph(
        "A balanced squad requires adequate depth in every position group. "
        "The heatmap below highlights which clubs are well-covered and which have critical gaps."
    )
    if os.path.exists(f"{IMG_DIR}/ro2_position_heatmap.png"):
        doc.add_picture(f"{IMG_DIR}/ro2_position_heatmap.png", width=Inches(6))

    if not ro_players.empty:
        doc.add_heading("24.2.1 Position Gap Diagnosis", level=3)
        gap_notes = []
        for team in sorted(ro_players["short_team"].unique()):
            tp = ro_players[ro_players["short_team"] == team]
            pos_counts = tp["pos_group"].value_counts().to_dict()
            gaps = []
            for pos, ideal in IDEAL_SQUAD.items():
                actual = pos_counts.get(pos, 0)
                if actual < ideal:
                    gaps.append(f"{pos} (have {actual}, need {ideal})")
            if gaps:
                gap_notes.append(f"{team}: Under-covered in {', '.join(gaps)}")
            else:
                gap_notes.append(f"{team}: Fully covered across all positions")
        for note in gap_notes:
            p = doc.add_paragraph(note)
            p.paragraph_format.space_after = Pt(2)

    # 24.3 Nationality Balance
    doc.add_heading("24.3 Nationality Balance & Foreign Talent", level=2)
    doc.add_paragraph(
        "Romanian Liga I clubs balance between developing local talent and importing foreign players. "
        "The ratio impacts squad cohesion, development pathways, and compliance with domestic player rules."
    )
    if os.path.exists(f"{IMG_DIR}/ro3_foreign_domestic.png"):
        doc.add_picture(f"{IMG_DIR}/ro3_foreign_domestic.png", width=Inches(6))

    if not ro_players.empty and "strNationality" in ro_players.columns:
        doc.add_heading("24.3.1 Top Nationalities in Liga I", level=3)
        nat = ro_players["strNationality"].value_counts().head(15)
        rows = [[n, c, f"{c / len(ro_players) * 100:.1f}%"] for n, c in nat.items()]
        _add_table(doc, ["Nationality", "Players", "Share"], rows)

    # =====================================================================
    # SECTION 25: PLAYER SCOUTING - BEST PLAYERS & HIDDEN GEMS
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("25. Player Scouting - Best Players & Hidden Gems", level=1)

    doc.add_paragraph(
        "Using market value data, age profiles, and positional analysis, we identify the most valuable "
        "players, best young prospects, and undervalued gems that represent transfer opportunities."
    )

    # 25.1 Top Players by Market Value
    doc.add_heading("25.1 Top 20 Most Valuable Players", level=2)
    if os.path.exists(f"{IMG_DIR}/ro4_top_value_players.png"):
        doc.add_picture(f"{IMG_DIR}/ro4_top_value_players.png", width=Inches(6))

    if not ro_players.empty:
        top20 = ro_players.nlargest(20, "market_value_eur_m")
        rows = []
        for i, (_, p) in enumerate(top20.iterrows(), 1):
            rows.append([i, p["strPlayer"], p["short_team"], p.get("strPosition", "?"),
                         int(p["age"]) if pd.notna(p["age"]) else "?",
                         p.get("strNationality", "?"),
                         f"€{p['market_value_eur_m']:.2f}M"])
        _add_table(doc, ["#", "Player", "Club", "Position", "Age", "Nationality", "Value"], rows)

    # 25.2 Best Players by Position
    doc.add_heading("25.2 Best Players by Position", level=2)
    doc.add_paragraph(
        "Identifying the best player in each position group helps scouts target specific needs."
    )
    if not ro_players.empty:
        for pos_full, pos_short in sorted(POS_GROUPS.items(), key=lambda x: x[1]):
            pos_players = ro_players[ro_players["strPosition"] == pos_full].nlargest(3, "market_value_eur_m")
            if not pos_players.empty and pos_players["market_value_eur_m"].sum() > 0:
                doc.add_heading(f"{pos_full}", level=3)
                rows = []
                for _, p in pos_players.iterrows():
                    rows.append([p["strPlayer"], p["short_team"],
                                 int(p["age"]) if pd.notna(p["age"]) else "?",
                                 p.get("strNationality", "?"),
                                 f"€{p['market_value_eur_m']:.2f}M"])
                _add_table(doc, ["Player", "Club", "Age", "Nationality", "Value"], rows)
                doc.add_paragraph("")

    # 25.3 Wonderkids & Young Prospects
    doc.add_heading("25.3 Wonderkids & Young Prospects (U21)", level=2)
    doc.add_paragraph(
        "Players under 21 represent the future of Romanian football. Those with high market values "
        "at a young age are prime targets for bigger leagues and could generate significant transfer fees."
    )
    if not ro_players.empty:
        u21 = ro_players[(ro_players["age"] <= 21) & (ro_players["age"] > 16)].nlargest(15, "market_value_eur_m")
        if not u21.empty:
            rows = []
            for _, p in u21.iterrows():
                rows.append([p["strPlayer"], p["short_team"], p.get("strPosition", "?"),
                             int(p["age"]), p.get("strNationality", "?"),
                             f"€{p['market_value_eur_m']:.2f}M"])
            _add_table(doc, ["Player", "Club", "Position", "Age", "Nationality", "Value"], rows)

    # 25.4 Undervalued Gems - High Performing Low Market Value
    doc.add_heading("25.4 Undervalued Gems - Experienced Players at Bargain Prices", level=2)
    doc.add_paragraph(
        "Experienced players (25-29 years old) with relatively low market values can provide "
        "immediate quality at minimal cost - ideal for clubs with tight budgets."
    )
    if not ro_players.empty:
        bargains = ro_players[
            (ro_players["age"] >= 25) & (ro_players["age"] <= 29) &
            (ro_players["market_value_eur_m"] > 0) & (ro_players["market_value_eur_m"] <= 0.5)
        ].sort_values("market_value_eur_m", ascending=False).head(15)
        if not bargains.empty:
            rows = []
            for _, p in bargains.iterrows():
                rows.append([p["strPlayer"], p["short_team"], p.get("strPosition", "?"),
                             int(p["age"]), f"€{p['market_value_eur_m']:.2f}M"])
            _add_table(doc, ["Player", "Club", "Position", "Age", "Value"], rows)

    # =====================================================================
    # SECTION 26: CLUB-BY-CLUB DIAGNOSTICS
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("26. Club-by-Club Diagnostics", level=1)

    doc.add_paragraph(
        "Each club receives a diagnostic report covering performance metrics, squad composition, "
        "strengths, weaknesses, and recommended actions. This section uses all available data "
        "to provide actionable insights for management and coaching staff."
    )

    if not ro_standings.empty:
        completed = pd.DataFrame()
        if not ro_events.empty:
            completed = ro_events.dropna(subset=["intHomeScore", "intAwayScore"]).copy()

        for _, team_row in ro_standings.iterrows():
            team_full = team_row["strTeam"]
            team_short = SHORT_NAMES.get(team_full, team_full)

            doc.add_heading(f"26.{int(team_row['intRank'])} {team_full}", level=2)

            # Performance summary
            group = "Playoff" if team_row["intRank"] <= 6 else "Playout"
            doc.add_heading("Performance Summary", level=3)
            perf_rows = [
                ["Position", f"#{int(team_row['intRank'])} ({group})"],
                ["Points (Current)", int(team_row["intPoints"])],
                ["Played", int(team_row["intPlayed"])],
                ["Record", f"{int(team_row['intWin'])}W {int(team_row['intDraw'])}D {int(team_row['intLoss'])}L"],
                ["Goals For", int(team_row["intGoalsFor"])],
                ["Goals Against", int(team_row["intGoalsAgainst"])],
                ["Goal Difference", int(team_row["intGoalDifference"])],
                ["Form (Last 5)", team_row.get("strForm", "N/A")],
                ["PPG", f"{team_row['intPoints'] / max(team_row['intPlayed'], 1):.2f}"],
            ]
            if "intRegSeasonPts" in team_row.index:
                perf_rows.append(["Regular Season Pts", int(team_row["intRegSeasonPts"])])
            if "intHalvedPts" in team_row.index:
                perf_rows.append(["Halved Pts (Split)", int(team_row["intHalvedPts"])])
            if "intPostSplitPts" in team_row.index:
                perf_rows.append(["Post-Split Pts", int(team_row["intPostSplitPts"])])
            _add_table(doc, ["Metric", "Value"], perf_rows)
            doc.add_paragraph("")

            # Home/Away split
            if not completed.empty:
                doc.add_heading("Home / Away Split", level=3)
                home_m = completed[completed["strHomeTeam"] == team_full]
                away_m = completed[completed["strAwayTeam"] == team_full]
                hw = len(home_m[home_m["intHomeScore"] > home_m["intAwayScore"]])
                hd = len(home_m[home_m["intHomeScore"] == home_m["intAwayScore"]])
                hl = len(home_m[home_m["intHomeScore"] < home_m["intAwayScore"]])
                hgf = int(home_m["intHomeScore"].sum())
                hga = int(home_m["intAwayScore"].sum())
                aw = len(away_m[away_m["intAwayScore"] > away_m["intHomeScore"]])
                ad = len(away_m[away_m["intAwayScore"] == away_m["intHomeScore"]])
                al = len(away_m[away_m["intAwayScore"] < away_m["intHomeScore"]])
                agf = int(away_m["intAwayScore"].sum())
                aga = int(away_m["intHomeScore"].sum())
                _add_table(doc, ["", "W", "D", "L", "GF", "GA", "GD"],
                           [["Home", hw, hd, hl, hgf, hga, hgf - hga],
                            ["Away", aw, ad, al, agf, aga, agf - aga]])
                doc.add_paragraph("")

                # Key insight
                if hw + hd + hl > 0 and aw + ad + al > 0:
                    home_ppg = (hw * 3 + hd) / (hw + hd + hl)
                    away_ppg = (aw * 3 + ad) / (aw + ad + al)
                    if home_ppg > away_ppg * 1.5:
                        doc.add_paragraph(
                            f"⚠️ {team_short} shows a severe home/away imbalance "
                            f"(Home PPG: {home_ppg:.2f} vs Away PPG: {away_ppg:.2f}). "
                            "Improving away performance should be a priority."
                        )
                    elif away_ppg > home_ppg * 1.3:
                        doc.add_paragraph(
                            f"💡 {team_short} performs better away ({away_ppg:.2f} PPG) "
                            f"than at home ({home_ppg:.2f} PPG) - an unusual pattern that "
                            "suggests potential stadium/fan pressure issues."
                        )

            # Squad composition
            tp = ro_players[ro_players["short_team"] == team_short]
            if tp.empty:
                # Try broader match
                tp = ro_players[ro_players["strTeam"].str.contains(team_full.split()[0], case=False, na=False)]

            if not tp.empty:
                doc.add_heading("Squad Composition", level=3)
                avg_age = tp["age"].mean()
                total_val = tp["market_value_eur_m"].sum()
                u21_count = (tp["age"] <= 21).sum()
                foreign_count = (~tp["strNationality"].str.contains("Romania", case=False, na=False)).sum() if "strNationality" in tp.columns else 0
                squad_rows = [
                    ["Squad Size", len(tp)],
                    ["Average Age", f"{avg_age:.1f}"],
                    ["Total Market Value", f"€{total_val:.2f}M"],
                    ["U21 Players", u21_count],
                    ["Foreign Players", foreign_count],
                    ["Romanian Players", len(tp) - foreign_count],
                ]
                _add_table(doc, ["Metric", "Value"], squad_rows)
                doc.add_paragraph("")

                # Strengths & Weaknesses
                doc.add_heading("Strengths & Weaknesses", level=3)
                strengths = []
                weaknesses = []

                # Scoring analysis
                gf = int(team_row["intGoalsFor"])
                ga = int(team_row["intGoalsAgainst"])
                avg_gf = ro_standings["intGoalsFor"].mean()
                avg_ga = ro_standings["intGoalsAgainst"].mean()

                if gf > avg_gf * 1.15:
                    strengths.append(f"Strong attack ({gf} goals, {((gf / avg_gf - 1) * 100):.0f}% above average)")
                elif gf < avg_gf * 0.85:
                    weaknesses.append(f"Weak attack ({gf} goals, {((1 - gf / avg_gf) * 100):.0f}% below average)")

                if ga < avg_ga * 0.85:
                    strengths.append(f"Solid defense (only {ga} conceded, {((1 - ga / avg_ga) * 100):.0f}% below average)")
                elif ga > avg_ga * 1.15:
                    weaknesses.append(f"Leaky defense ({ga} conceded, {((ga / avg_ga - 1) * 100):.0f}% above average)")

                if avg_age < 25:
                    strengths.append(f"Young squad (avg {avg_age:.1f}) with high development potential")
                elif avg_age > 29:
                    weaknesses.append(f"Aging squad (avg {avg_age:.1f}) requiring renewal")

                if u21_count >= 5:
                    strengths.append(f"Strong youth pipeline ({u21_count} U21 players)")
                elif u21_count <= 2:
                    weaknesses.append(f"Poor youth integration (only {u21_count} U21 players)")

                if total_val > 15:
                    strengths.append(f"High squad value (€{total_val:.1f}M) indicates quality depth")
                elif total_val < 5:
                    weaknesses.append(f"Low squad value (€{total_val:.1f}M) limits competitiveness")

                # Position gaps
                pos_counts = tp["pos_group"].value_counts().to_dict()
                for pos, ideal in IDEAL_SQUAD.items():
                    if pos_counts.get(pos, 0) < ideal:
                        weaknesses.append(f"Thin at {pos} ({pos_counts.get(pos, 0)}/{ideal})")

                if team_row.get("strForm", "").count("W") >= 3:
                    strengths.append(f"Good current form ({team_row.get('strForm', 'N/A')})")
                elif team_row.get("strForm", "").count("L") >= 3:
                    weaknesses.append(f"Poor current form ({team_row.get('strForm', 'N/A')})")

                if strengths:
                    doc.add_paragraph("Strengths:", style="List Bullet")
                    for s in strengths:
                        p = doc.add_paragraph(s, style="List Bullet 2")
                if weaknesses:
                    doc.add_paragraph("Weaknesses:", style="List Bullet")
                    for w in weaknesses:
                        p = doc.add_paragraph(w, style="List Bullet 2")

            doc.add_paragraph("")

    # =====================================================================
    # SECTION 27: TRANSFER STRATEGY & RECOMMENDATIONS
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("27. Transfer Strategy & Recommendations", level=1)

    doc.add_paragraph(
        "Based on the squad diagnostics above, this section provides specific transfer "
        "recommendations for each club, identifying positions that need reinforcement and "
        "potential targets from within the league or abroad."
    )

    # 27.1 Strength Radar
    doc.add_heading("27.1 Top 6 Clubs - Strength Radar Comparison", level=2)
    doc.add_paragraph(
        "The radar chart compares the top 6 clubs across five key dimensions: squad depth, "
        "youth pipeline, goal difference strength, market value, and foreign talent integration. "
        "Areas where a club falls short indicate priority transfer needs."
    )
    if os.path.exists(f"{IMG_DIR}/ro8_club_radar.png"):
        doc.add_picture(f"{IMG_DIR}/ro8_club_radar.png", width=Inches(5.5))

    # 27.2 Club Market Value
    doc.add_heading("27.2 Squad Value by Position", level=2)
    doc.add_paragraph(
        "Understanding how each club invests across positions reveals strategic priorities. "
        "Clubs with low investment in attacking positions despite poor goal records need "
        "offensive reinforcements."
    )
    if os.path.exists(f"{IMG_DIR}/ro6_squad_value_composition.png"):
        doc.add_picture(f"{IMG_DIR}/ro6_squad_value_composition.png", width=Inches(6))

    # 27.3 Transfer Recommendations per Club
    doc.add_heading("27.3 Transfer Recommendations per Club", level=2)

    if not ro_standings.empty and not ro_players.empty:
        for _, team_row in ro_standings.iterrows():
            team_full = team_row["strTeam"]
            team_short = SHORT_NAMES.get(team_full, team_full)
            tp = ro_players[ro_players["short_team"] == team_short]
            if tp.empty:
                tp = ro_players[ro_players["strTeam"].str.contains(team_full.split()[0], case=False, na=False)]
            if tp.empty:
                continue

            doc.add_heading(f"{team_short}", level=3)

            # Determine needs
            needs = []
            pos_counts = tp["pos_group"].value_counts().to_dict()
            for pos, ideal in IDEAL_SQUAD.items():
                actual = pos_counts.get(pos, 0)
                if actual < ideal:
                    needs.append(f"Sign {ideal - actual} {pos}")

            avg_age = tp["age"].mean()
            if avg_age > 28:
                needs.append("Prioritize younger signings (avg age too high)")
            
            gf = int(team_row["intGoalsFor"])
            ga = int(team_row["intGoalsAgainst"])
            avg_gf = ro_standings["intGoalsFor"].mean()
            avg_ga = ro_standings["intGoalsAgainst"].mean()
            if gf < avg_gf:
                needs.append("Recruit attacking talent (below average scoring)")
            if ga > avg_ga:
                needs.append("Strengthen defense (above average goals conceded)")

            # Players to sell (high value, over 30)
            sellable = tp[(tp["age"] >= 31) & (tp["market_value_eur_m"] > 0.3)].nlargest(
                3, "market_value_eur_m"
            )

            if needs:
                doc.add_paragraph("Transfer Needs:")
                for n in needs:
                    doc.add_paragraph(n, style="List Bullet 2")

            if not sellable.empty:
                doc.add_paragraph("Potential Sales (30+, generate funds):")
                for _, p in sellable.iterrows():
                    doc.add_paragraph(
                        f"{p['strPlayer']} ({p.get('strPosition', '?')}, age {int(p['age'])}, "
                        f"€{p['market_value_eur_m']:.2f}M)",
                        style="List Bullet 2"
                    )

            # Suggested targets from within the league
            if needs:
                target_positions = []
                for n in needs:
                    for pos_short in IDEAL_SQUAD:
                        if pos_short in n:
                            target_positions.append(pos_short)
                if target_positions:
                    targets = ro_players[
                        (ro_players["short_team"] != team_short) &
                        (ro_players["pos_group"].isin(target_positions)) &
                        (ro_players["age"] <= 27) &
                        (ro_players["market_value_eur_m"] > 0)
                    ].nlargest(3, "market_value_eur_m")
                    if not targets.empty:
                        doc.add_paragraph("Internal Transfer Targets:")
                        for _, t in targets.iterrows():
                            doc.add_paragraph(
                                f"{t['strPlayer']} ({t.get('strPosition', '?')}, {t['short_team']}, "
                                f"age {int(t['age'])}, €{t['market_value_eur_m']:.2f}M)",
                                style="List Bullet 2"
                            )
            doc.add_paragraph("")

    # =====================================================================
    # SECTION 28: VIDEO / MATCH EVENT ANALYSIS
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("28. Match Event Analysis & Key Moments", level=1)

    doc.add_paragraph(
        "While full video analysis requires match footage, our event data captures every match result, "
        "score, date, venue, and round. This section reconstructs key narrative moments of the season - "
        "biggest wins, heaviest defeats, high-scoring thrillers, and decisive matches."
    )

    # 28.1 Home vs Away Butterfly
    doc.add_heading("28.1 Home vs Away Goal Efficiency", level=2)
    doc.add_paragraph(
        "The butterfly chart shows each club's home and away goal output side-by-side, "
        "revealing which teams are fortress-like at home and which perform better on the road."
    )
    if os.path.exists(f"{IMG_DIR}/ro5_home_away_butterfly.png"):
        doc.add_picture(f"{IMG_DIR}/ro5_home_away_butterfly.png", width=Inches(6))

    # 28.2 Season Form Heatmap
    doc.add_heading("28.2 Season Form Heatmap", level=2)
    doc.add_paragraph(
        "The round-by-round heatmap below shows every team's result in every round. "
        "Green cells represent wins (3 points), orange for draws (1 point), and red for losses (0). "
        "This visualization reveals momentum streaks, collapses, and recovery patterns."
    )
    if os.path.exists(f"{IMG_DIR}/ro7_form_heatmap.png"):
        doc.add_picture(f"{IMG_DIR}/ro7_form_heatmap.png", width=Inches(6.5))

    # 28.3 Biggest Results of the Season
    doc.add_heading("28.3 Biggest Results of the Season", level=2)
    if not ro_events.empty:
        completed = ro_events.dropna(subset=["intHomeScore", "intAwayScore"]).copy()
        if not completed.empty:
            completed["total_goals"] = completed["intHomeScore"] + completed["intAwayScore"]
            completed["margin"] = abs(completed["intHomeScore"] - completed["intAwayScore"])

            # Highest scoring matches
            doc.add_heading("Highest Scoring Matches", level=3)
            top_goals = completed.nlargest(10, "total_goals")
            rows = []
            for _, m in top_goals.iterrows():
                rows.append([
                    m.get("dateEvent", "").strftime("%Y-%m-%d") if hasattr(m.get("dateEvent", ""), "strftime") else str(m.get("dateEvent", ""))[:10],
                    f"{SHORT_NAMES.get(m['strHomeTeam'], m['strHomeTeam'])} {int(m['intHomeScore'])} - {int(m['intAwayScore'])} {SHORT_NAMES.get(m['strAwayTeam'], m['strAwayTeam'])}",
                    int(m["total_goals"]),
                    m.get("strVenue", "N/A")
                ])
            _add_table(doc, ["Date", "Match", "Goals", "Venue"], rows)

            # Biggest upsets (biggest margin wins by lower-ranked teams)
            doc.add_paragraph("")
            doc.add_heading("Biggest Margin Victories", level=3)
            top_margin = completed.nlargest(10, "margin")
            rows = []
            for _, m in top_margin.iterrows():
                winner = m["strHomeTeam"] if m["intHomeScore"] > m["intAwayScore"] else m["strAwayTeam"]
                rows.append([
                    m.get("dateEvent", "").strftime("%Y-%m-%d") if hasattr(m.get("dateEvent", ""), "strftime") else str(m.get("dateEvent", ""))[:10],
                    f"{SHORT_NAMES.get(m['strHomeTeam'], m['strHomeTeam'])} {int(m['intHomeScore'])} - {int(m['intAwayScore'])} {SHORT_NAMES.get(m['strAwayTeam'], m['strAwayTeam'])}",
                    int(m["margin"]),
                    SHORT_NAMES.get(winner, winner)
                ])
            _add_table(doc, ["Date", "Match", "Margin", "Winner"], rows)

            # Draw analysis
            doc.add_paragraph("")
            doc.add_heading("Draw Analysis", level=3)
            draws = completed[completed["intHomeScore"] == completed["intAwayScore"]]
            doc.add_paragraph(
                f"Total draws this season: {len(draws)} out of {len(completed)} matches "
                f"({len(draws) / len(completed) * 100:.1f}%). "
                f"Zero-zero draws: {len(draws[draws['intHomeScore'] == 0])}."
            )

            # Monthly goal trends
            if "dateEvent" in completed.columns:
                doc.add_heading("Monthly Scoring Trends", level=3)
                completed["month"] = completed["dateEvent"].dt.to_period("M")
                monthly = completed.groupby("month").agg(
                    Matches=("total_goals", "count"),
                    TotalGoals=("total_goals", "sum"),
                ).reset_index()
                monthly["AvgGoals"] = (monthly["TotalGoals"] / monthly["Matches"]).round(2)
                rows = [[str(r["month"]), int(r["Matches"]), int(r["TotalGoals"]), r["AvgGoals"]]
                        for _, r in monthly.iterrows()]
                _add_table(doc, ["Month", "Matches", "Total Goals", "Avg Goals/Match"], rows)

    # =====================================================================
    # SECTION 29: SEASON NARRATIVE & PREDICTIONS
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("29. Season Narrative & Predictions", level=1)

    doc.add_paragraph(
        "Drawing together all the data points from the preceding analysis, this final section "
        "constructs the overall narrative of the 2025-2026 Romanian Liga I season and offers "
        "data-driven predictions for the remaining matches."
    )

    # 29.1 Title Race
    doc.add_heading("29.1 Title Race Analysis", level=2)
    if not ro_standings.empty:
        playoff = ro_standings[ro_standings["intRank"] <= 6].copy()
        if not playoff.empty:
            leader = playoff.iloc[0]
            second = playoff.iloc[1] if len(playoff) > 1 else None
            doc.add_paragraph(
                f"The title race is led by {leader['strTeam']} with {int(leader['intPoints'])} points. "
            )
            if second is not None:
                gap = int(leader["intPoints"]) - int(second["intPoints"])
                doc.add_paragraph(
                    f"{second['strTeam']} sits {gap} points behind in 2nd place with "
                    f"{int(second['intPoints'])} points. "
                    f"{'The gap is virtually insurmountable.' if gap > 8 else 'The race remains open.' if gap <= 4 else 'The leader has a comfortable but not decisive advantage.'}"
                )

            doc.add_paragraph("")
            doc.add_heading("Playoff Table", level=3)
            rows = []
            for _, t in playoff.iterrows():
                form = t.get("strForm", "N/A")
                rows.append([int(t["intRank"]), t["strTeam"], int(t["intPoints"]),
                             int(t["intPlayed"]),
                             f"{int(t['intWin'])}W {int(t['intDraw'])}D {int(t['intLoss'])}L",
                             int(t["intGoalDifference"]), form])
            _add_table(doc, ["#", "Team", "Pts", "P", "Record", "GD", "Form"], rows)

    # 29.2 Relegation Battle
    doc.add_heading("29.2 Relegation Battle", level=2)
    if not ro_standings.empty:
        playout = ro_standings[ro_standings["intRank"] > 6].copy()
        if not playout.empty:
            bottom2 = playout.tail(2)
            doc.add_paragraph(
                f"The bottom of the table sees {bottom2.iloc[-1]['strTeam']} "
                f"({int(bottom2.iloc[-1]['intPoints'])} pts) in last place, "
                f"{'already looking certain for relegation' if int(bottom2.iloc[-1]['intPoints']) < int(playout.iloc[0]['intPoints']) - 15 else 'still fighting to survive'}. "
            )
            if len(bottom2) > 1:
                doc.add_paragraph(
                    f"{bottom2.iloc[0]['strTeam']} ({int(bottom2.iloc[0]['intPoints'])} pts) "
                    "also faces serious danger."
                )

            doc.add_paragraph("")
            doc.add_heading("Playout Table", level=3)
            rows = []
            for _, t in playout.iterrows():
                rows.append([int(t["intRank"]), t["strTeam"], int(t["intPoints"]),
                             int(t["intPlayed"]),
                             f"{int(t['intWin'])}W {int(t['intDraw'])}D {int(t['intLoss'])}L",
                             int(t["intGoalDifference"]), t.get("strForm", "N/A")])
            _add_table(doc, ["#", "Team", "Pts", "P", "Record", "GD", "Form"], rows)

    # 29.3 Key Takeaways
    doc.add_heading("29.3 Key Takeaways", level=2)
    takeaways = []
    if not ro_standings.empty:
        best_attack = ro_standings.nlargest(1, "intGoalsFor").iloc[0]
        best_defense = ro_standings.nsmallest(1, "intGoalsAgainst").iloc[0]
        worst_attack = ro_standings.nsmallest(1, "intGoalsFor").iloc[0]
        worst_defense = ro_standings.nlargest(1, "intGoalsAgainst").iloc[0]
        takeaways.append(
            f"Best attack: {best_attack['strTeam']} ({int(best_attack['intGoalsFor'])} goals)"
        )
        takeaways.append(
            f"Best defense: {best_defense['strTeam']} ({int(best_defense['intGoalsAgainst'])} conceded)"
        )
        takeaways.append(
            f"Weakest attack: {worst_attack['strTeam']} ({int(worst_attack['intGoalsFor'])} goals)"
        )
        takeaways.append(
            f"Leakiest defense: {worst_defense['strTeam']} ({int(worst_defense['intGoalsAgainst'])} conceded)"
        )

        # PPG leader
        ro_standings["ppg_calc"] = ro_standings["intPoints"] / ro_standings["intPlayed"].replace(0, 1)
        ppg_leader = ro_standings.nlargest(1, "ppg_calc").iloc[0]
        takeaways.append(
            f"Highest PPG: {ppg_leader['strTeam']} ({ppg_leader['ppg_calc']:.2f})"
        )

    if not ro_players.empty:
        most_valuable_team = ro_players.groupby("short_team")["market_value_eur_m"].sum().idxmax()
        mv_total = ro_players.groupby("short_team")["market_value_eur_m"].sum().max()
        takeaways.append(f"Most valuable squad: {most_valuable_team} (€{mv_total:.1f}M)")

        youngest_team = ro_players.groupby("short_team")["age"].mean().idxmin()
        youngest_age = ro_players.groupby("short_team")["age"].mean().min()
        takeaways.append(f"Youngest squad: {youngest_team} (avg {youngest_age:.1f} years)")

    for t in takeaways:
        doc.add_paragraph(t, style="List Bullet")

    # 29.4 Predictions
    doc.add_heading("29.4 Data-Driven Predictions", level=2)
    if not ro_standings.empty:
        predictions = []
        leader = ro_standings.iloc[0]
        predictions.append(
            f"Champion: {leader['strTeam']} — currently {int(leader['intPoints'])} points with "
            f"form '{leader.get('strForm', 'N/A')}', holding the strongest combined metrics."
        )
        bottom = ro_standings.iloc[-1]
        predictions.append(
            f"Relegated: {bottom['strTeam']} — {int(bottom['intPoints'])} points, "
            f"goal difference of {int(bottom['intGoalDifference'])}, mathematical survival increasingly unlikely."
        )
        # FCSB analysis (dropped to playout despite being reigning champion)
        fcsb = ro_standings[ro_standings["strTeam"].str.contains("FCSB", case=False)]
        if not fcsb.empty:
            fcsb_row = fcsb.iloc[0]
            if int(fcsb_row["intRank"]) > 6:
                predictions.append(
                    f"FCSB Alert: The reigning champions sit in the Playout group "
                    f"(#{int(fcsb_row['intRank'])}, {int(fcsb_row['intPoints'])} pts). "
                    "Despite having the league's most valuable squad, inconsistency and tactical "
                    "issues have derailed their title defense."
                )

        for pred in predictions:
            doc.add_paragraph(pred, style="List Bullet")

    return doc
