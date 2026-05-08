"""
Report Document Part 1 - Word document sections 1-8
Covers: Title, TOC, Executive Summary, Objective, Data Sources, Methodology,
        League Overview, Match Analysis, Championship Probability, European Comparison
"""

import os, glob
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from report_data import IMG_DIR


def build_document_part1(doc, data):
    """Build the first half of the Word document (sections 1-8)."""
    standings = data["standings"]
    events = data["events"]
    players = data["players"]
    ro_standings = data["ro_standings"]
    ro_events = data["ro_events"]
    playoff_teams_data = data["playoff_teams_data"]
    playout_teams_data = data["playout_teams_data"]

    # ===================== TITLE PAGE =====================
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Football Analytics Pro")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Complete Data Analysis\nRomanian Superliga 2025-2026")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Final Project - Skillab Sports Analytics Course\nLesson 11: Complete Data Analysis for a Team or League")
    run.font.size = Pt(13)
    run.font.italic = True

    doc.add_paragraph()

    stats_p = doc.add_paragraph()
    stats_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stats_p.add_run(f"Coverage: 15 European Leagues | {len(standings)} team records | {len(events)} match events | {len(players)} players")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Author: Adrian Dome\nDate: May 2026\nPlatform: Streamlit Cloud (Python)\nGitHub: github.com/domeadrian/football-analytics")
    run.font.size = Pt(11)

    doc.add_page_break()

    # ===================== TABLE OF CONTENTS =====================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Objective & Hypothesis",
        "3. Data Sources & Quality",
        "4. Methodology & KPIs",
        "5. League Overview & Standings Analysis",
        "6. Match Analysis & Results",
        "7. Championship Probability (Monte Carlo)",
        "8. European Benchmarking",
        "9. Player Analysis & Database",
        "10. Scouting & Recruitment Intelligence",
        "11. Tactical Analysis",
        "12. Opponent Report & Team Comparison",
        "13. Transfer Recommendations",
        "14. Physical & Medical Analysis",
        "15. Youth Academy & Development",
        "16. Financial Analysis",
        "17. Season Projections",
        "18. Machine Learning Models",
        "19. Advanced Statistics",
        "20. League Management & Venues",
        "21. Dashboard Platform & Technical Stack",
        "22. Conclusions & Impact",
        "23. Technical Appendix & References",
        "",
        "--- Romanian Liga I Deep Dive ---",
        "24. Squad Analysis Deep Dive",
        "25. Player Scouting - Best Players & Hidden Gems",
        "26. Club-by-Club Diagnostics",
        "27. Transfer Strategy & Recommendations",
        "28. Match Event Analysis & Key Moments",
        "29. Season Narrative & Predictions",
    ]
    for item in toc_items:
        if item.startswith("---") or item == "":
            doc.add_paragraph(item)
        else:
            p = doc.add_paragraph(item, style="List Number")
            p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ===================== 1. EXECUTIVE SUMMARY =====================
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This report presents a comprehensive, professional-grade data analysis of the Romanian "
        "Superliga (Liga I) 2025-2026 season. Built as an interactive analytics platform covering "
        "15 European leagues, it demonstrates the full spectrum of modern football analytics: "
        "from descriptive statistics and match analysis to predictive machine learning models "
        "and prescriptive transfer recommendations."
    )
    doc.add_paragraph(
        "The project is deployed as a live Streamlit web application with 19 specialized modules, "
        "each addressing a different aspect of football club operations. The platform processes "
        f"{len(standings)} team standings records, {len(events)} match events, and {len(players)} "
        "player profiles across all 15 leagues."
    )

    doc.add_heading("Key Findings", level=2)
    findings = [
        "The Romanian Superliga playoff/playout system creates distinct competitive tiers that can be modeled statistically",
        "Monte Carlo simulation (10,000 iterations) provides robust championship probability estimates for the 6 playoff teams",
        "Machine Learning clustering (K-Means) successfully identifies 4 performance tiers that align with league position",
        "xPoints analysis reveals which teams are over/underperforming relative to expected output",
        "Cross-league financial analysis shows Romanian clubs achieve higher points-per-value efficiency than top-5 league clubs",
    ]
    if not ro_standings.empty:
        leader = ro_standings.iloc[0]["strTeam"]
        leader_pts = ro_standings.iloc[0]["intPoints"]
        findings.insert(1, f"Current playoff leader: {leader} with {leader_pts} points")
    for f in findings:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("Report Scope", level=2)
    scope_table = doc.add_table(rows=1, cols=3)
    scope_table.style = 'Light Shading Accent 1'
    hdr = scope_table.rows[0].cells
    hdr[0].text = "Dimension"
    hdr[1].text = "Coverage"
    hdr[2].text = "Detail"
    scope_data = [
        ("Leagues", "15", "Top European leagues + Romanian Liga I & II"),
        ("Teams", str(len(standings)), "Full standings with playoff/playout split"),
        ("Matches", str(len(events)), "All 2025-2026 season events with scores"),
        ("Players", str(len(players)), "Full squad data with ages, positions, nationalities"),
        ("Charts", "35+", "Plotly visualizations across all analysis modules"),
        ("ML Models", "4", "K-Means, Random Forest, Linear Regression, Poisson"),
        ("Dashboard Pages", "19", "Full interactive Streamlit platform"),
    ]
    for dim, cov, det in scope_data:
        row = scope_table.add_row().cells
        row[0].text = dim
        row[1].text = cov
        row[2].text = det

    doc.add_page_break()

    # ===================== 2. OBJECTIVE & HYPOTHESIS =====================
    doc.add_heading("2. Objective & Hypothesis", level=1)

    doc.add_heading("2.1 Project Objective", level=2)
    doc.add_paragraph(
        "To build a complete, professional-grade football analytics platform that answers:"
    )
    p = doc.add_paragraph()
    run = p.add_run(
        '"Can data-driven analysis of the Romanian Superliga identify which teams are over/underperforming '
        'relative to their expected output, and can this be used to predict championship outcomes '
        'and inform transfer decisions?"'
    )
    run.font.italic = True

    doc.add_heading("2.2 Hypothesis", level=2)
    doc.add_paragraph(
        "H1: Teams with positive xG (expected goals) differential and high PPG efficiency correlate "
        "with higher championship probability in Monte Carlo simulation."
    )
    doc.add_paragraph(
        "H2: Statistical models (K-Means clustering, Random Forest) can identify team performance tiers "
        "that predict end-of-season positions better than simple points-based ranking."
    )
    doc.add_paragraph(
        "H3: Cross-league scouting using normalized metrics can identify undervalued players suitable "
        "for recruitment by Romanian Superliga clubs."
    )

    doc.add_heading("2.3 Key Performance Indicators", level=2)
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.style = 'Light Shading Accent 1'
    hdr = kpi_table.rows[0].cells
    hdr[0].text = "KPI"
    hdr[1].text = "Formula"
    hdr[2].text = "Purpose"
    hdr[3].text = "Target"
    kpis = [
        ("PPG", "Points / Played", "Efficiency measure", "> 2.0 for title"),
        ("Goal Difference", "GF - GA", "Net performance", "> +15 for top 3"),
        ("xG (Poisson)", "λ = Att×Avg", "Expected scoring", "Match prediction"),
        ("xPoints", "P(W)×3 + P(D)×1", "Expected points", "Luck detection"),
        ("Win Rate %", "Wins / Played × 100", "Form indicator", "> 50% elite"),
        ("Championship Prob.", "MC sim wins / 10K", "Title likelihood", "Probabilistic"),
        ("ELO Rating", "1500 + GD×5 + Pts×2", "Power ranking", "> 1600 strong"),
        ("Consistency (H)", "-Σ p·log₂(p)", "Result entropy", "< 1.2 predictable"),
    ]
    for kpi, formula, purpose, target in kpis:
        row = kpi_table.add_row().cells
        row[0].text = kpi
        row[1].text = formula
        row[2].text = purpose
        row[3].text = target

    doc.add_page_break()

    # ===================== 3. DATA SOURCES =====================
    doc.add_heading("3. Data Sources & Quality", level=1)

    doc.add_heading("3.1 Primary Data Sources", level=2)
    ds_table = doc.add_table(rows=1, cols=5)
    ds_table.style = 'Light Shading Accent 1'
    hdr = ds_table.rows[0].cells
    hdr[0].text = "Source"
    hdr[1].text = "Data Type"
    hdr[2].text = "Coverage"
    hdr[3].text = "Method"
    hdr[4].text = "License"
    sources = [
        ("TheSportsDB", "Standings, events, teams, players", "15 leagues, 2025-26", "REST API", "CC BY-NC-SA 4.0"),
        ("Football-Data.org", "Competition metadata", "European leagues", "REST API", "Free Tier"),
        ("Transfermarkt", "Market values, ages, contracts", "All scope teams", "Web scraping", "Fair use"),
    ]
    for src, dtype, cov, method, lic in sources:
        row = ds_table.add_row().cells
        row[0].text = src
        row[1].text = dtype
        row[2].text = cov
        row[3].text = method
        row[4].text = lic

    doc.add_heading("3.2 Data Volume & Storage", level=2)
    doc.add_paragraph(f"• Standings records: {len(standings)} (15 leagues)")
    doc.add_paragraph(f"• Match events: {len(events)} (all completed matches)")
    doc.add_paragraph(f"• Player profiles: {len(players)} (ages, positions, nationalities)")
    doc.add_paragraph(f"• Romanian Liga I: {len(ro_standings)} teams, {len(ro_events)} match events")
    doc.add_paragraph(f"• Data files: {len(glob.glob('data/*.json'))} JSON files in data/ directory")
    doc.add_paragraph(f"• Total storage: ~5 MB structured JSON data")

    doc.add_heading("3.3 Data Quality Pipeline", level=2)
    doc.add_paragraph("The ETL (Extract-Transform-Load) pipeline implements:")
    quality_steps = [
        "Type coercion: pd.to_numeric(errors='coerce') for all numeric fields",
        "Missing value imputation: NaN → 0 for numeric, NaN → 'Unknown' for categorical",
        "Date parsing: ISO 8601 format with timezone awareness",
        "Deduplication: Primary key validation across multiple source files",
        "Cross-reference: Standings ↔ Events consistency check (team names, rounds played)",
        "League ID mapping: Consistent naming across 3 different data sources",
        "Split-league handling: Points halving and group assignment for Romania/Belgium/Scotland",
        "Automated updates: auto_update_data.py script fetches fresh data daily",
    ]
    for step in quality_steps:
        doc.add_paragraph(step, style="List Bullet")

    doc.add_page_break()

    # ===================== 4. METHODOLOGY =====================
    doc.add_heading("4. Methodology & Analytical Framework", level=1)

    doc.add_heading("4.1 Statistical Methods", level=2)
    methods_table = doc.add_table(rows=1, cols=4)
    methods_table.style = 'Light Shading Accent 1'
    hdr = methods_table.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "Application"
    hdr[2].text = "Parameters"
    hdr[3].text = "Reference"
    methods = [
        ("Poisson Distribution", "xG & match prediction", "λ = team attack × league avg", "Dixon & Coles (1997)"),
        ("Monte Carlo", "Championship probability", "N=10,000 iterations", "Metropolis & Ulam (1949)"),
        ("ELO Rating", "Power ranking", "K=32, Base=1500", "Hvattum & Arntzen (2010)"),
        ("K-Means Clustering", "Team tier classification", "k=4, StandardScaler", "MacQueen (1967)"),
        ("Random Forest", "Feature importance", "n_estimators=100", "Breiman (2001)"),
        ("Linear Regression", "Points prediction", "OLS, R² validation", "Standard OLS"),
        ("Shannon Entropy", "Consistency index", "H = -Σ p·log₂(p)", "Shannon (1948)"),
    ]
    for method, app, params, ref in methods:
        row = methods_table.add_row().cells
        row[0].text = method
        row[1].text = app
        row[2].text = params
        row[3].text = ref

    doc.add_heading("4.2 Analysis Framework (4-Level)", level=2)
    doc.add_paragraph("The analysis follows a structured 4-level framework:")
    levels = [
        ("Descriptive", "What happened?", "Standings, results, goal patterns, form sequences"),
        ("Diagnostic", "Why?", "xG analysis, home/away splits, attack/defense quadrants"),
        ("Predictive", "What next?", "Monte Carlo, ML models, season projections, ELO"),
        ("Prescriptive", "What to do?", "Scouting recommendations, transfer targets, tactical adjustments"),
    ]
    level_table = doc.add_table(rows=1, cols=4)
    level_table.style = 'Light Shading Accent 1'
    hdr = level_table.rows[0].cells
    hdr[0].text = "Level"
    hdr[1].text = "Question"
    hdr[2].text = "Techniques Used"
    hdr[3].text = " "
    for lvl, q, tech in levels:
        row = level_table.add_row().cells
        row[0].text = lvl
        row[1].text = q
        row[2].text = tech

    doc.add_heading("4.3 Romanian Liga I Format", level=2)
    doc.add_paragraph(
        "The Romanian Superliga uses a unique playoff/playout format that requires special modeling:"
    )
    format_steps = [
        "Regular Season: 16 teams play 30 rounds (each team plays all others twice)",
        "Split Point: After Round 30, league splits into two groups",
        "Championship Playoff (Top 6): Points are HALVED (rounded up), then 10 more rounds played",
        "Relegation Playout (Bottom 10): 18 more rounds; last 2 relegated, 13th-14th play relegation playoff",
        "Title Decision: Only playoff teams can win the championship; winner gets UCL spot",
        "European Spots: 1st → UCL, 2nd → UEL, Cup winner or 3rd → UECL",
    ]
    for step in format_steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_page_break()

    # ===================== 5. LEAGUE OVERVIEW =====================
    doc.add_heading("5. League Overview & Standings Analysis", level=1)

    doc.add_heading("5.1 Full League Points Distribution", level=2)
    doc.add_paragraph(
        "The bar chart below shows all 16 Romanian Liga I teams ranked by current total points. "
        "The split between playoff (top 6) and playout (bottom 10) is clearly visible in the "
        "points gap around the 6th-7th position boundary."
    )
    if os.path.exists(f"{IMG_DIR}/01_standings_full.png"):
        doc.add_picture(f"{IMG_DIR}/01_standings_full.png", width=Inches(6))

    doc.add_heading("5.2 Championship Playoff Standings", level=2)
    doc.add_paragraph(
        "After the split, the top 6 teams had their regular season points halved before starting "
        "the playoff phase. This compression makes the playoff highly competitive."
    )
    if os.path.exists(f"{IMG_DIR}/01_playoff_standings.png"):
        doc.add_picture(f"{IMG_DIR}/01_playoff_standings.png", width=Inches(6))

    # Insert playoff table
    if playoff_teams_data:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Championship Playoff - Detailed Statistics").bold = True
        st_table = doc.add_table(rows=1, cols=8)
        st_table.style = 'Light Shading Accent 1'
        hdr = st_table.rows[0].cells
        for i, h in enumerate(["Pos", "Team", "P", "W", "D", "L", "GD", "Pts"]):
            hdr[i].text = h
        for idx, t in enumerate(sorted(playoff_teams_data, key=lambda x: -x["intPoints"]), 1):
            cells = st_table.add_row().cells
            cells[0].text = str(idx)
            cells[1].text = str(t["strTeam"])
            cells[2].text = str(t["intPlayed"])
            cells[3].text = str(t["intWin"])
            cells[4].text = str(t["intDraw"])
            cells[5].text = str(t["intLoss"])
            cells[6].text = str(t["intGoalDifference"])
            cells[7].text = str(t["intPoints"])

    doc.add_heading("5.3 Relegation Playout - Danger Zones", level=2)
    doc.add_paragraph(
        "The bottom 10 teams compete in the Relegation Playout. Teams are color-coded by zone: "
        "green = safe, orange = relegation playoff (13th-14th play vs Liga II 3rd-4th), "
        "red = direct relegation (15th-16th)."
    )
    if os.path.exists(f"{IMG_DIR}/01_playout_standings.png"):
        doc.add_picture(f"{IMG_DIR}/01_playout_standings.png", width=Inches(6))

    # Insert playout table
    if playout_teams_data:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Relegation Playout - Detailed Statistics").bold = True
        st_table2 = doc.add_table(rows=1, cols=8)
        st_table2.style = 'Light Shading Accent 1'
        hdr = st_table2.rows[0].cells
        for i, h in enumerate(["Pos", "Team", "P", "W", "D", "L", "GD", "Pts"]):
            hdr[i].text = h
        for idx, t in enumerate(sorted(playout_teams_data, key=lambda x: -x["intPoints"]), 7):
            cells = st_table2.add_row().cells
            cells[0].text = str(idx)
            cells[1].text = str(t["strTeam"])
            cells[2].text = str(t["intPlayed"])
            cells[3].text = str(t["intWin"])
            cells[4].text = str(t["intDraw"])
            cells[5].text = str(t["intLoss"])
            cells[6].text = str(t["intGoalDifference"])
            cells[7].text = str(t["intPoints"])

    doc.add_heading("5.4 Cross-League Comparison - Top 3 per League", level=2)
    doc.add_paragraph(
        "The following chart shows the top 3 teams from all 15 leagues side by side, "
        "providing context for where Romanian teams stand in the European landscape."
    )
    if os.path.exists(f"{IMG_DIR}/01_all_leagues_top3.png"):
        doc.add_picture(f"{IMG_DIR}/01_all_leagues_top3.png", width=Inches(6))

    doc.add_page_break()

    # ===================== 6. MATCH ANALYSIS =====================
    doc.add_heading("6. Match Analysis & Results", level=1)

    doc.add_heading("6.1 Match Outcome Distribution", level=2)
    doc.add_paragraph(
        "The pie chart shows the proportion of home wins, draws, and away wins across all "
        "Romanian Liga I matches this season. This reveals the strength of home advantage."
    )
    if os.path.exists(f"{IMG_DIR}/02_outcome_distribution.png"):
        doc.add_picture(f"{IMG_DIR}/02_outcome_distribution.png", width=Inches(4.5))

    # Calculate and add stats
    if not ro_events.empty and "intHomeScore" in ro_events.columns:
        completed = ro_events.dropna(subset=["intHomeScore", "intAwayScore"])
        if not completed.empty:
            total_m = len(completed)
            doc.add_paragraph(f"Total completed matches analyzed: {total_m}")
            avg_goals = (completed["intHomeScore"].astype(float) + completed["intAwayScore"].astype(float)).mean()
            doc.add_paragraph(f"Average goals per match: {avg_goals:.2f}")

    doc.add_heading("6.2 Goals per Match Distribution", level=2)
    doc.add_paragraph(
        "This histogram shows how many goals are typically scored per match. The distribution "
        "helps identify whether the league is high-scoring or defensive, and informs Poisson "
        "model parameters."
    )
    if os.path.exists(f"{IMG_DIR}/02_goals_distribution.png"):
        doc.add_picture(f"{IMG_DIR}/02_goals_distribution.png", width=Inches(5.5))

    doc.add_heading("6.3 Scoreline Heatmap", level=2)
    doc.add_paragraph(
        "The heatmap reveals the most common scorelines. The concentration around 1-0, 1-1, "
        "and 2-1 results is typical of European football, with diminishing frequency for "
        "high-scoring matches."
    )
    if os.path.exists(f"{IMG_DIR}/02_score_heatmap.png"):
        doc.add_picture(f"{IMG_DIR}/02_score_heatmap.png", width=Inches(4.5))

    doc.add_heading("6.4 Monthly Goals Trend", level=2)
    doc.add_paragraph(
        "The time series analysis shows how goal-scoring patterns evolve throughout the season. "
        "Early-season matches often have higher goal counts as teams are still finding their "
        "defensive shape, while late-season fixtures tend to be tighter."
    )
    if os.path.exists(f"{IMG_DIR}/02_monthly_trend.png"):
        doc.add_picture(f"{IMG_DIR}/02_monthly_trend.png", width=Inches(6))

    doc.add_page_break()

    # ===================== 7. CHAMPIONSHIP PROBABILITY =====================
    doc.add_heading("7. Championship Probability (Monte Carlo Simulation)", level=1)

    doc.add_heading("7.1 Methodology", level=2)
    doc.add_paragraph(
        "A Monte Carlo simulation with 10,000 iterations was used to estimate championship "
        "probabilities. The simulation works as follows:"
    )
    mc_steps = [
        "For each of the 6 playoff teams, calculate current win rate (W%), draw rate (D%), loss rate (L%)",
        "For each remaining match, generate a random outcome using these rates as probabilities",
        "Sum the current points + simulated points to get final projected total",
        "Record which team finishes with the most points (random tiebreaker for ties)",
        "After 10,000 iterations, the championship probability = times finished 1st / 10,000",
    ]
    for step in mc_steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Important: ").bold = True
    p.add_run("Only the 6 playoff teams are simulated, as the championship is decided exclusively "
              "within this group. Playout teams cannot win the title by definition of the format.")

    doc.add_heading("7.2 Championship Probability Results", level=2)
    if os.path.exists(f"{IMG_DIR}/03_monte_carlo.png"):
        doc.add_picture(f"{IMG_DIR}/03_monte_carlo.png", width=Inches(6))

    doc.add_heading("7.3 Points Distribution Range (Box Plot)", level=2)
    doc.add_paragraph(
        "The box plot shows the range of simulated final points for each team across all "
        "10,000 iterations. The overlap between boxes indicates uncertainty in final positions."
    )
    if os.path.exists(f"{IMG_DIR}/03_points_range.png"):
        doc.add_picture(f"{IMG_DIR}/03_points_range.png", width=Inches(6))

    doc.add_heading("7.4 Top 3 Contenders - Win Probability Gauges", level=2)
    doc.add_paragraph(
        "The gauge indicators provide an at-a-glance view of the top 3 championship contenders, "
        "showing their probability of winning the title based on current form."
    )
    if os.path.exists(f"{IMG_DIR}/03_gauges.png"):
        doc.add_picture(f"{IMG_DIR}/03_gauges.png", width=Inches(6))

    doc.add_page_break()

    # ===================== 8. EUROPEAN COMPARISON =====================
    doc.add_heading("8. European Benchmarking", level=1)

    doc.add_heading("8.1 Goals Scored Comparison", level=2)
    doc.add_paragraph(
        "Comparing average goals scored per team across all 15 leagues reveals which leagues "
        "are more attacking. The color represents league strength coefficient."
    )
    if os.path.exists(f"{IMG_DIR}/04_euro_goals.png"):
        doc.add_picture(f"{IMG_DIR}/04_euro_goals.png", width=Inches(6))

    doc.add_heading("8.2 Competitiveness Index", level=2)
    doc.add_paragraph(
        "The competitiveness index uses points standard deviation within each league. A lower "
        "value means teams are closer together in quality, making outcomes less predictable. "
        "Leagues with playoff systems (Romania, Belgium) often show compressed competitiveness."
    )
    if os.path.exists(f"{IMG_DIR}/04_competitiveness.png"):
        doc.add_picture(f"{IMG_DIR}/04_competitiveness.png", width=Inches(6))

    doc.add_heading("8.3 League Strength vs Market Value", level=2)
    doc.add_paragraph(
        "The bubble chart maps league strength (UEFA coefficient proxy) against total market "
        "value. Bubble size represents total goals scored. This reveals the relationship "
        "between financial power and competitive quality."
    )
    if os.path.exists(f"{IMG_DIR}/04_strength_market.png"):
        doc.add_picture(f"{IMG_DIR}/04_strength_market.png", width=Inches(6))

    doc.add_heading("8.4 Romanian Liga I in European Context", level=2)
    doc.add_paragraph("Key comparative insights:")
    euro_insights = [
        "Romania's league strength coefficient (0.35) places it in the lower half of European football",
        "However, the playoff system creates elite-level competitiveness among the top 6 teams",
        "Market value (€8M avg) is significantly lower than top-5 leagues (€280-500M), but higher than Scandinavia",
        "Goals-per-team ratio is comparable to mid-tier leagues like Turkey and Greece",
        "The playoff compression effect means Romanian top teams compete at a higher relative intensity",
    ]
    for insight in euro_insights:
        doc.add_paragraph(insight, style="List Bullet")

    doc.add_page_break()

    return doc
